import csv
import re
import xml.etree.ElementTree as ET
import os
import json
import sys
from docx import Document
from openpyxl import load_workbook
import pdfplumber
import chardet

#GASs
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'src')))
from storage.mongo import save_to_mongo 

def extract_articles_fields(article):
    return {
        "source_name": article.get("source_name") or article.get("source", {}).get("name"),
        "title": article.get("title"),
        "author": article.get("author"),
        "description": article.get("description"),
        "publishedAt": article.get("publishedAt"),
        "content": article.get("content")
    }

def parse_csv_file(file_path):
    with open(file_path, "r") as f:
        reader = csv.DictReader(f)  
        for row in reader:
            # print(f"Source name: {row['source_name']}, Title: {row['title']}, Author: {row['author']}, Description: {row['description']}, Published At: {row['publishedAt']}, Content: {row['content']}")
            article_fields = extract_articles_fields(row)  
            print(f"Saving to MongoDB: {article_fields}")
            save_to_mongo(article_fields, "CSV Source")
            

def parse_xml_file(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    for item in root.findall("article"):
        article = {
            "source_name": item.find("source/name").text,
            "title": item.find("title").text,
            "author": item.find("author").text,
            "description": item.find("description").text,
            "publishedAt": item.find("publishedAt").text,
            "content": item.find("content").text
        }
        # print(f"Source Name: {article['source_name']}, Title: {article['title']}, Author: {article['author']}, Description: {article['description']}, Published At: {article['publishedAt']}, Content: {article['content']}")
        article_fields = extract_articles_fields(article) 
        print(f"Saving to MongoDB: {article_fields}")
        save_to_mongo(article_fields, "XML Source")

def parse_json_files():
    folder_path = "../../data/raw/api/"

    for filename in os.listdir(folder_path):
        if filename.endswith(".json"): 
            file_path = os.path.join(folder_path, filename)

            with open(file_path, "r") as f:
                article_data = json.load(f)

            if isinstance(article_data, dict):  # If article_data is a dictionary (single article)
                article_fields = extract_articles_fields(article_data)
                # print(article_fields)
                print(f"Saving to MongoDB: {article_fields}")
                save_to_mongo(article_fields, filename)
            else:
                print(f"Unexpected structure in {filename}: {article_data}")

def extract_text_from_pdf(pdf_path):
    pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(normalize_text(text))

    return "\n\n".join(pages)

def normalize_text(text):
    if not text:
        return ""

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()


def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    confidence = result.get("confidence")

    print(f"Detected encoding: {encoding} (confidence: {confidence})")

    text = raw.decode(encoding, errors="replace")

    return text

def extract_text_from_two_column_pdf(pdf_path, gap=10):
    pages_text = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            mid_x = page.width / 2

            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))

            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")

            combined = "\n\n".join(part for part in [left_text, right_text] if part)
            if combined:
                pages_text.append(combined)

    return "\n\n".join(pages_text)

def extract_text_from_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

def extract_text_from_two_column_word(docx_path):
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

def extract_data_from_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb["Articles"]

    movies = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        movie = {
            "ID": row[0],
            "Title": row[1],
            "Source": row[2],
            "Author": row[3],
            "Published Date": row[4],
            "Word Count": row[5],
            "Category": row[6]
        }
        movies.append(movie)

    return movies
def extract_summary_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Category Summary"]

    summary = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        metric = row[0]
        value = row[1]
        if metric is None:
            continue
        summary[str(metric)] = value

    return summary

if __name__ == "__main__":
    # Assignment 3
    # parse_json_files()
    # parse_csv_file("../../data/raw/csv/articles.csv")
    # parse_xml_file("../../data/raw/xml/articles.xml")

    #Assignment 4

    # pdf_path = "../../data/raw/research/health_digest_normal.pdf"
    # text = extract_text_from_pdf(pdf_path)
    # print(text)
    
    # pdf_path2 = "../../data/raw/research/health_digest_two_column.pdf"
    # text2 = extract_text_from_two_column_pdf(pdf_path2)
    # print(text2)

    # docx_path = "../../data/raw/docx/health_digest_normal.docx"
    # text = extract_text_from_word(docx_path)
    # print(text)

    excel_path = "../../data/raw/xlsx/health_articles.xlsx"
    articles = extract_data_from_excel(excel_path)
    print(articles)
